from __future__ import annotations

from pathlib import Path
from datetime import datetime

import pandas as pd
from docx import Document

APP_DIR = Path(__file__).resolve().parents[1]
ASSETS_DIR = APP_DIR / "assets"
DATA_DIR = APP_DIR / "data"
DOCX_DIR = DATA_DIR / "generated_docs"
DOCX_DIR.mkdir(parents=True, exist_ok=True)

TEMPLATE_CANDIDATES = [
    ASSETS_DIR / "skierowanie_template.docx",
    ASSETS_DIR / "Skierowanie na badania lekarskie.docx",
    APP_DIR / "skierowanie_template.docx",
    APP_DIR / "Skierowanie na badania lekarskie.docx",
]

CATEGORY_TO_SECTION = {
    "CZYNNIKI FIZYCZNE": "I",
    "PYŁY": "II",
    "CZYNNIKI CHEMICZNE": "III",
    "CZYNNIKI BIOLOGICZNE": "IV",
    "INNE": "V",
}


def _find_template() -> Path | None:
    for path in TEMPLATE_CANDIDATES:
        if path.exists():
            return path
    return None


def _fmt_date(value) -> str:
    if value is None or value == "":
        return ""
    if isinstance(value, str):
        if "T" in value:
            value = value.split("T")[0]
        for fmt in ("%Y-%m-%d", "%d.%m.%Y"):
            try:
                return datetime.strptime(value, fmt).strftime("%d.%m.%Y")
            except Exception:
                pass
        try:
            return datetime.fromisoformat(value).strftime("%d.%m.%Y")
        except Exception:
            return value
    try:
        return value.strftime("%d.%m.%Y")
    except Exception:
        return str(value)


def _extract_hazards(referral: dict) -> pd.DataFrame:
    raw = referral.get("hazards")

    if isinstance(raw, pd.DataFrame):
        df = raw.copy()
    elif isinstance(raw, list):
        df = pd.DataFrame(raw)
    else:
        df = pd.DataFrame()

    if df.empty:
        return pd.DataFrame(columns=["hazard_name", "category", "section_label", "work_conditions"])

    rename_map = {
        "Zagrożenie": "hazard_name",
        "hazard": "hazard_name",
        "Kategoria": "category",
        "Sekcja": "section_label",
        "Opis warunków pracy": "work_conditions",
    }
    df = df.rename(columns=rename_map)

    for col in ["hazard_name", "category", "section_label", "work_conditions"]:
        if col not in df.columns:
            df[col] = ""

    df["hazard_name"] = df["hazard_name"].fillna("").astype(str).str.strip()
    df["category"] = df["category"].fillna("").astype(str).str.strip().str.upper()
    df["section_label"] = df["section_label"].fillna("").astype(str).str.strip()
    df["work_conditions"] = df["work_conditions"].fillna("").astype(str).str.strip()

    df = df[df["hazard_name"] != ""].copy()

    if df.empty:
        return df

    missing = df["section_label"] == ""
    df.loc[missing, "section_label"] = df.loc[missing, "category"].map(CATEGORY_TO_SECTION).fillna("V")

    df["section_label"] = df["section_label"].replace(
        {
            "I. Czynniki fizyczne": "I",
            "II. Pyły": "II",
            "III. Czynniki chemiczne": "III",
            "IV. Czynniki biologiczne": "IV",
            "V. Inne czynniki, w tym niebezpieczne": "V",
        }
    )

    return df


def _group_hazards_by_section(referral: dict) -> dict[str, str]:
    df = _extract_hazards(referral)
    result = {k: "" for k in ["I", "II", "III", "IV", "V"]}

    if df.empty:
        return result

    for section in result:
        vals = df[df["section_label"] == section]["hazard_name"].tolist()
        result[section] = "; ".join(v for v in vals if str(v).strip())

    return result


def strike(text: str) -> dict[str, object]:
    return {"text": text, "strike": True}


def normal(text: str) -> dict[str, object]:
    return {"text": text, "strike": False}


def _replace_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return

    new_text = full_text
    changed = False

    for key, value in mapping.items():
        if key in new_text:
            new_text = new_text.replace(key, value)
            changed = True

    if not changed:
        return

    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new_text


def _replace_everywhere(doc: Document, mapping: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, mapping)


def _replace_exam_type_placeholders(doc: Document, exam_type: str) -> None:
    exam_type = (exam_type or "").strip().lower()

    values = {
        "wstepne": normal("wstępne") if exam_type == "wstępne" else strike("wstępne"),
        "okresowe": normal("okresowe") if exam_type == "okresowe" else strike("okresowe"),
        "kontrolne": normal("kontrolne") if exam_type == "kontrolne" else strike("kontrolne"),
    }

    def replace_in_paragraph(paragraph):
        tokens = ["{{wstepne}}", "{{okresowe}}", "{{kontrolne}}"]
        text = "".join(run.text for run in paragraph.runs)
        if not text:
            return
        if not any(token in text for token in tokens):
            return

        parts = []
        i = 0
        while i < len(text):
            matched = False
            for token, value in [
                ("{{wstepne}}", values["wstepne"]),
                ("{{okresowe}}", values["okresowe"]),
                ("{{kontrolne}}", values["kontrolne"]),
            ]:
                if text.startswith(token, i):
                    parts.append(value)
                    i += len(token)
                    matched = True
                    break
            if not matched:
                parts.append({"text": text[i], "strike": False})
                i += 1

        for run in paragraph.runs:
            run.text = ""

        first = True
        for part in parts:
            if first and paragraph.runs:
                run = paragraph.runs[0]
                first = False
            else:
                run = paragraph.add_run()
            run.text = part["text"]
            run.font.strike = part["strike"]

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)


def generate_referral_docx(referral: dict) -> str:
    template_path = _find_template()
    if not template_path:
        raise FileNotFoundError(
            "Nie znaleziono szablonu Word. Umieść plik w assets/skierowanie_template.docx"
        )

    referral_number = str(
        referral.get("referral_number", f"ref_{datetime.now().strftime('%Y%m%d%H%M%S')}")
    )
    safe_name = referral_number.replace("/", "-").replace("\\", "-")
    out_path = DOCX_DIR / f"{safe_name}.docx"

    employer = str(referral.get("employer", "") or "SAFETY Service")
    place_of_issue = str(referral.get("place_of_issue", "") or "Warszawa")
    issue_date = _fmt_date(referral.get("issue_date"))
    employee_name = str(referral.get("employee_name", "") or "")
    pesel = str(referral.get("pesel", "") or "")
    employee_address = str(referral.get("employee_address", "") or "")
    position_name = str(referral.get("position_name", "") or "")
    position_description = str(referral.get("position_description", "") or position_name)
    exam_type = str(referral.get("exam_type", "") or "")

    grouped = _group_hazards_by_section(referral)
    hazards_df = _extract_hazards(referral)
    hazards_count = len(hazards_df)

    doc = Document(str(template_path))

    mapping = {
        "{{employer}}": employer,
        "{{place_and_date}}": f"{place_of_issue} {issue_date}".strip(),
        "{{employee_name}}": employee_name,
        "{{pesel}}": pesel,
        "{{address}}": employee_address,
        "{{position}}": position_name,
        "{{position_description}}": position_description,
        "{{I}}": grouped["I"] or "-",
        "{{II}}": grouped["II"] or "-",
        "{{III}}": grouped["III"] or "-",
        "{{IV}}": grouped["IV"] or "-",
        "{{V}}": grouped["V"] or "-",
        "{{hazards_count}}": str(hazards_count),
    }

    _replace_everywhere(doc, mapping)
    _replace_exam_type_placeholders(doc, exam_type)

    doc.save(str(out_path))
    return str(out_path)